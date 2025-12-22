"""Document Processing Service for RAG knowledge base.

This service handles:
- PDF loading and text extraction (PyPDFLoader)
- DOCX loading with section/heading detection (python-docx)
- Document chunking with overlap
- Metadata enrichment (section tracking, tenant isolation)
- Integration with LangChain document loaders

Supported formats:
- .pdf: Page-based extraction
- .docx/.doc: Paragraph-based with heading/section tracking

"""
from typing import List, Dict, Any, Optional
from datetime import datetime
import pytz
from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from src.utils.logging import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Service for processing documents for RAG ingestion."""

    def __init__(
        self,
        chunk_size: int = 900,
        chunk_overlap: int = 150,
        separators: Optional[List[str]] = None
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Maximum characters per chunk (default: 600)
            chunk_overlap: Character overlap between chunks (default: 200)
            separators: Custom split separators (default: paragraph/sentence/word)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Default separators optimized for technical documentation
        self.separators = separators or [
            "\n\n",  # Paragraphs
            "\n",    # Lines
            ". ",    # Sentences
            " ",     # Words
            ""       # Characters (fallback)
        ]

        # Initialize text splitter
        self._init_splitter()

        logger.info(
            "document_processor_initialized",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            overlap_percentage=f"{(chunk_overlap/chunk_size)*100:.1f}%"
        )

    def _init_splitter(self):
        """Initialize text splitter with current config."""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=self.separators,
            is_separator_regex=False
        )

    def update_config(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
        separators: Optional[List[str]] = None
    ):
        """
        Update processor config dynamically.

        Args:
            chunk_size: New chunk size (optional)
            chunk_overlap: New chunk overlap (optional)
            separators: New separators (optional)
        """
        if chunk_size is not None:
            self.chunk_size = chunk_size
        if chunk_overlap is not None:
            self.chunk_overlap = chunk_overlap
        if separators is not None:
            self.separators = separators

        # Reinitialize splitter with new config
        self._init_splitter()

        logger.info(
            "document_processor_config_updated",
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            overlap_percentage=f"{(self.chunk_overlap/self.chunk_size)*100:.1f}%"
        )

    def load_pdf(self, pdf_path: str) -> List[Document]:
        """
        Load PDF and return LangChain documents with cleaned text.

        Improvements:
        - Removes image-related artifacts and captions
        - Cleans up excessive whitespace
        - Preserves page structure for better chunking

        Args:
            pdf_path: Path to PDF file

        Returns:
            List of Document objects with page_content and metadata

        Metadata includes:
            - source: PDF file path
            - page: Page number (0-indexed)
        """
        try:
            # Validate file exists
            if not Path(pdf_path).exists():
                raise FileNotFoundError(f"PDF file not found: {pdf_path}")

            logger.info("loading_pdf", pdf_path=pdf_path)

            # Load PDF using PyPDFLoader
            loader = PyPDFLoader(pdf_path)
            raw_documents = loader.load()

            import re

            # Clean up each page
            documents = []
            for doc in raw_documents:
                text = doc.page_content
                
                # Remove common image artifacts
                # Remove lines that are likely image captions
                lines = text.split('\n')
                cleaned_lines = []
                
                for line in lines:
                    line_lower = line.strip().lower()
                    
                    # Skip image-related lines
                    if any(indicator in line_lower for indicator in [
                        'figure ', 'image ', 'img ', 'photo ', 'picture ', 
                        'diagram ', 'screenshot', 'illustration'
                    ]):
                        # Only skip if it's a short line (likely a caption)
                        if len(line.strip()) < 100:
                            continue
                    
                    # Skip very short lines that might be image titles
                    if len(line.strip()) > 0 and len(line.strip()) < 5:
                        continue
                    
                    cleaned_lines.append(line)
                
                # Rejoin and clean up whitespace
                cleaned_text = '\n'.join(cleaned_lines)
                
                # Remove excessive newlines (more than 2 consecutive)
                cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
                
                # Remove excessive spaces
                cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)
                
                # Only add if there's meaningful content
                if cleaned_text.strip():
                    doc.page_content = cleaned_text.strip()
                    documents.append(doc)

            logger.info(
                "pdf_loaded_successfully",
                pdf_path=pdf_path,
                page_count=len(documents),
                total_chars=sum(len(doc.page_content) for doc in documents),
                avg_chars_per_page=sum(len(doc.page_content) for doc in documents) / len(documents) if documents else 0
            )

            return documents

        except Exception as e:
            logger.error(
                "pdf_load_failed",
                pdf_path=pdf_path,
                error=str(e)
            )
            raise

    def load_docx(self, docx_path: str) -> List[Document]:
        """
        Load DOCX and return LangChain documents with section tracking.

        Improved approach:
        - Combines paragraphs within sections to create larger text blocks
        - Extracts table content as formatted text
        - Skips image captions and titles
        - Prevents creation of thousands of tiny documents

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of Document objects with page_content and metadata

        Metadata includes:
            - source: DOCX file path
            - file_type: '.docx'
            - section_title: Current section heading
            - section_number: Extracted section number (e.g., "2.3.3")
            - has_tables: Whether this section contains tables
        """
        try:
            # Validate file exists
            if not Path(docx_path).exists():
                raise FileNotFoundError(f"DOCX file not found: {docx_path}")

            logger.info("loading_docx", docx_path=docx_path)

            # Import python-docx
            try:
                from docx import Document as DocxDocument
                from docx.oxml.text.paragraph import CT_P
                from docx.oxml.table import CT_Tbl
                from docx.table import _Cell, Table
                from docx.text.paragraph import Paragraph
            except ImportError:
                raise ImportError(
                    "python-docx is required for DOCX processing. "
                    "Install with: pip install python-docx>=1.1.0"
                )

            import re

            # Load DOCX
            doc = DocxDocument(docx_path)

            # Helper function to extract table as text
            def extract_table_text(table: Table) -> str:
                """Extract table content as formatted text."""
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                return "\n".join(table_text)

            # Helper function to check if paragraph is likely an image caption/title
            def is_image_caption(para) -> bool:
                """Check if paragraph is likely an image caption or title."""
                text = para.text.strip().lower()
                style = para.style.name.lower()
                
                # Skip if it's a caption style
                if 'caption' in style or 'figure' in style:
                    return True
                
                # Skip if text starts with common image indicators
                image_indicators = ['figure', 'image', 'img', 'photo', 'picture', 'diagram']
                if any(text.startswith(ind) for ind in image_indicators):
                    return True
                
                # Skip very short paragraphs that might be image titles (< 10 chars)
                if len(text) < 10 and len(text) > 0:
                    return True
                
                return False

            # Process document by iterating through all elements (paragraphs and tables)
            documents = []
            current_section = None
            current_section_number = None
            section_content = []
            section_has_tables = False

            def create_section_document():
                """Create a document from accumulated section content."""
                if section_content:
                    combined_text = "\n\n".join(section_content)
                    if combined_text.strip():
                        document = Document(
                            page_content=combined_text,
                            metadata={
                                'source': docx_path,
                                'file_type': '.docx',
                                'section_title': current_section or 'Unknown',
                                'section_number': current_section_number,
                                'has_tables': section_has_tables
                            }
                        )
                        documents.append(document)
                        logger.debug(
                            "section_document_created",
                            section_title=current_section,
                            content_length=len(combined_text),
                            has_tables=section_has_tables
                        )

            # Iterate through document body elements (paragraphs and tables)
            for element in doc.element.body:
                # Check if it's a paragraph
                if isinstance(element, CT_P):
                    para = Paragraph(element, doc)
                    text = para.text.strip()
                    
                    if not text:
                        continue
                    
                    # Skip image captions
                    if is_image_caption(para):
                        logger.debug("skipping_image_caption", text=text[:50])
                        continue
                    
                    is_heading = para.style.name.startswith("Heading")
                    
                    # If this is a heading, save previous section and start new one
                    if is_heading:
                        # Save previous section
                        create_section_document()
                        
                        # Start new section
                        current_section = text
                        match = re.match(r'^(\d+(?:\.\d+)*)\.\s+', text)
                        current_section_number = match.group(1) if match else None
                        section_content = []
                        section_has_tables = False
                        
                        logger.debug(
                            "new_section_started",
                            section_title=current_section,
                            section_number=current_section_number
                        )
                    else:
                        # Add paragraph to current section
                        section_content.append(text)
                
                # Check if it's a table
                elif isinstance(element, CT_Tbl):
                    table = Table(element, doc)
                    table_text = extract_table_text(table)
                    
                    if table_text.strip():
                        section_content.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
                        section_has_tables = True
                        logger.debug("table_extracted", rows=len(table.rows))

            # Don't forget the last section
            create_section_document()

            logger.info(
                "docx_loaded_successfully",
                docx_path=docx_path,
                section_count=len(documents),
                total_chars=sum(len(doc.page_content) for doc in documents),
                avg_chars_per_section=sum(len(doc.page_content) for doc in documents) / len(documents) if documents else 0
            )

            return documents

        except Exception as e:
            logger.error(
                "docx_load_failed",
                docx_path=docx_path,
                error=str(e)
            )
            raise

    def load_txt(self, txt_path: str) -> List[Document]:
        """
        Load TXT file and return LangChain documents (split by paragraphs).

        Useful for enriching knowledge base from chat history transcripts.

        Args:
            txt_path: Path to TXT file

        Returns:
            List of Document objects with page_content and metadata

        Metadata includes:
            - source: TXT file path
            - file_type: '.txt'
        """
        try:
            # Validate file exists
            if not Path(txt_path).exists():
                raise FileNotFoundError(f"TXT file not found: {txt_path}")

            logger.info("loading_txt", txt_path=txt_path)

            # Read text file
            with open(txt_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Split by double newlines (paragraphs)
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

            # Create LangChain documents (one per paragraph)
            documents = []
            for i, para in enumerate(paragraphs):
                document = Document(
                    page_content=para,
                    metadata={
                        'source': txt_path,
                        'file_type': '.txt',
                        'paragraph_index': i,
                    }
                )
                documents.append(document)

            logger.info(
                "txt_loaded_successfully",
                txt_path=txt_path,
                paragraph_count=len(documents),
                total_chars=sum(len(doc.page_content) for doc in documents)
            )

            return documents

        except Exception as e:
            logger.error(
                "txt_load_failed",
                txt_path=txt_path,
                error=str(e)
            )
            raise

    def chunk_documents(
        self,
        documents: List[Document],
        add_chunk_metadata: bool = True
    ) -> List[Document]:
        """
        Split documents into smaller chunks.

        Args:
            documents: List of LangChain Document objects
            add_chunk_metadata: Whether to add chunk_index to metadata

        Returns:
            List of chunked Document objects

        Metadata preservation:
            - Original metadata is preserved
            - New metadata added: chunk_index, chunk_total (if add_chunk_metadata=True)
        """
        try:
            logger.info(
                "chunking_documents",
                document_count=len(documents),
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )

            # Split documents
            chunks = self.text_splitter.split_documents(documents)

            # Add chunk metadata
            if add_chunk_metadata:
                chunk_index = 0
                for chunk in chunks:
                    chunk.metadata['chunk_index'] = chunk_index
                    chunk_index += 1

                # Add total chunk count to all chunks
                for chunk in chunks:
                    chunk.metadata['chunk_total'] = len(chunks)

            logger.info(
                "documents_chunked_successfully",
                original_document_count=len(documents),
                chunk_count=len(chunks),
                avg_chunk_size=sum(len(c.page_content) for c in chunks) / len(chunks) if chunks else 0
            )

            return chunks

        except Exception as e:
            logger.error(
                "document_chunking_failed",
                document_count=len(documents),
                error=str(e)
            )
            raise

    def enrich_metadata(
        self,
        documents: List[Document],
        tenant_id: str,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Add metadata to documents for multi-tenancy and tracking.

        Args:
            documents: List of Document objects
            tenant_id: Tenant UUID
            additional_metadata: Optional extra metadata to add

        Returns:
            Documents with enriched metadata

        Metadata added:
            - tenant_id: For multi-tenant isolation
            - ingested_at: ISO timestamp
            - Any additional_metadata provided
        """

        try:
            for doc in documents:
                # Add tenant_id (required for isolation)
                doc.metadata['tenant_id'] = tenant_id

                # Add timestamp
                doc.metadata['ingested_at'] = datetime.now(pytz.timezone('Asia/Ho_Chi_Minh')).isoformat()

                # Add any additional metadata
                if additional_metadata:
                    for key, value in additional_metadata.items():
                        # Don't override tenant_id
                        if key != 'tenant_id':
                            doc.metadata[key] = value

            logger.debug(
                "metadata_enriched",
                document_count=len(documents),
                tenant_id=tenant_id,
                additional_fields=list(additional_metadata.keys()) if additional_metadata else []
            )

            return documents

        except Exception as e:
            logger.error(
                "metadata_enrichment_failed",
                document_count=len(documents),
                error=str(e)
            )
            raise

    def process_pdf(
        self,
        pdf_path: str,
        tenant_id: str,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Complete PDF processing pipeline: Load → Chunk → Enrich.

        Args:
            pdf_path: Path to PDF file
            tenant_id: Tenant UUID
            additional_metadata: Optional metadata to add to all chunks

        Returns:
            List of processed Document chunks ready for embedding

        Pipeline:
            1. Load PDF (one Document per page)
            2. Chunk pages into smaller pieces
            3. Enrich metadata (tenant_id, timestamp, custom fields)
        """
        try:
            logger.info(
                "processing_pdf_started",
                pdf_path=pdf_path,
                tenant_id=tenant_id
            )

            # 1. Load PDF
            pages = self.load_pdf(pdf_path)

            # 2. Chunk documents
            chunks = self.chunk_documents(pages, add_chunk_metadata=True)

            # 3. Enrich metadata
            enriched_chunks = self.enrich_metadata(
                chunks,
                tenant_id=tenant_id,
                additional_metadata=additional_metadata
            )

            logger.info(
                "pdf_processing_completed",
                pdf_path=pdf_path,
                tenant_id=tenant_id,
                page_count=len(pages),
                chunk_count=len(enriched_chunks),
                avg_chars_per_chunk=sum(len(c.page_content) for c in enriched_chunks) / len(enriched_chunks)
            )

            return enriched_chunks

        except Exception as e:
            logger.error(
                "pdf_processing_failed",
                pdf_path=pdf_path,
                tenant_id=tenant_id,
                error=str(e)
            )
            raise

    def process_text(
        self,
        text: str,
        tenant_id: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Process raw text (alternative to PDF processing).

        Args:
            text: Raw text content
            tenant_id: Tenant UUID
            metadata: Optional metadata

        Returns:
            List of processed Document chunks
        """
        try:
            logger.info(
                "processing_text_started",
                text_length=len(text),
                tenant_id=tenant_id
            )

            # Create a single document from text
            doc = Document(
                page_content=text,
                metadata=metadata or {}
            )

            # Chunk it
            chunks = self.chunk_documents([doc], add_chunk_metadata=True)

            # Enrich metadata
            enriched_chunks = self.enrich_metadata(
                chunks,
                tenant_id=tenant_id,
                additional_metadata=metadata
            )

            logger.info(
                "text_processing_completed",
                text_length=len(text),
                tenant_id=tenant_id,
                chunk_count=len(enriched_chunks)
            )

            return enriched_chunks

        except Exception as e:
            logger.error(
                "text_processing_failed",
                text_length=len(text),
                tenant_id=tenant_id,
                error=str(e)
            )
            raise

    def process_document(
        self,
        file_path: str,
        tenant_id: str,
        additional_metadata: Optional[Dict[str, Any]] = None,
        chunk_config: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Universal document processing pipeline: Auto-detect format → Load → Chunk → Enrich.

        Supports:
            - .pdf: Uses PyPDFLoader
            - .docx/.doc: Uses python-docx with section tracking
            - .txt: Plain text split by paragraphs (useful for chat history enrichment)

        Args:
            file_path: Path to document file (.pdf, .docx, .doc, or .txt)
            tenant_id: Tenant UUID
            additional_metadata: Optional metadata to add to all chunks
            chunk_config: Optional chunking config (chunk_size, chunk_overlap, separators)

        Returns:
            List of processed Document chunks ready for embedding

        Pipeline:
            1. Detect file format by extension
            2. Load document (format-specific loader)
            3. Chunk documents into smaller pieces (with optional custom config)
            4. Enrich metadata (tenant_id, timestamp, custom fields)

        Example:
            >>> processor = get_document_processor()
            >>> chunks = processor.process_document(
            ...     file_path="eTMS.docx",
            ...     tenant_id="tenant-123"
            ... )
            >>> # Each chunk has section_title, section_number in metadata
        """
        try:
            file_path_obj = Path(file_path)
            file_ext = file_path_obj.suffix.lower()

            logger.info(
                "processing_document_started",
                file_path=file_path,
                file_type=file_ext,
                tenant_id=tenant_id
            )

            # Apply chunk configuration if provided
            original_chunk_size = self.chunk_size
            original_chunk_overlap = self.chunk_overlap
            original_separators = self.separators

            if chunk_config:
                # Temporarily update configuration
                self.update_config(
                    chunk_size=chunk_config.get('chunk_size'),
                    chunk_overlap=chunk_config.get('chunk_overlap'),
                    separators=chunk_config.get('separators')
                )

            try:
                # 1. Load based on format
                if file_ext == '.pdf':
                    documents = self.load_pdf(file_path)
                elif file_ext in ['.docx', '.doc']:
                    documents = self.load_docx(file_path)
                elif file_ext == '.txt':
                    documents = self.load_txt(file_path)
                else:
                    raise ValueError(
                        f"Unsupported file format: {file_ext}. "
                        f"Supported formats: .pdf, .docx, .doc, .txt"
                    )

                # 2. Chunk documents (preserves metadata including section_title)
                chunks = self.chunk_documents(documents, add_chunk_metadata=True)

                # 3. Enrich metadata with tenant info
                enriched_chunks = self.enrich_metadata(
                    chunks,
                    tenant_id=tenant_id,
                    additional_metadata=additional_metadata
                )

                logger.info(
                    "document_processing_completed",
                    file_path=file_path,
                    file_type=file_ext,
                    tenant_id=tenant_id,
                    original_document_count=len(documents),
                    chunk_count=len(enriched_chunks),
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap,
                    avg_chars_per_chunk=sum(len(c.page_content) for c in enriched_chunks) / len(enriched_chunks) if enriched_chunks else 0
                )

                return enriched_chunks

            finally:
                # Restore original configuration
                self.update_config(
                    chunk_size=original_chunk_size,
                    chunk_overlap=original_chunk_overlap,
                    separators=original_separators
                )

        except Exception as e:
            logger.error(
                "document_processing_failed",
                file_path=file_path,
                tenant_id=tenant_id,
                error=str(e)
            )
            raise


# Singleton instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor(
    chunk_size: int = 800,
    chunk_overlap: int = 200
) -> DocumentProcessor:
    """
    Get or create document processor singleton.

    Args:
        chunk_size: Chunk size (only used on first call)
        chunk_overlap: Chunk overlap (only used on first call)

    Returns:
        DocumentProcessor instance
    """
    global _document_processor

    if _document_processor is None:
        _document_processor = DocumentProcessor(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        logger.info("document_processor_singleton_created")

    return _document_processor
